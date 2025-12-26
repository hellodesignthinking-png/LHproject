"""
ZeroSite v4.0 LH Document Builder
=================================

python-docx 기반 LH 제안서 문서 생성

Author: ZeroSite M9 Team
Date: 2025-12-26
"""

import os
from typing import Dict, Any, List, Optional
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class LHDocumentBuilder:
    """LH 제안서 Word 문서 빌더"""
    
    def __init__(self):
        """문서 초기화"""
        self.doc = Document()
        self._setup_styles()
        
    def _setup_styles(self):
        """문서 스타일 설정"""
        # 제목 스타일
        styles = self.doc.styles
        
        # 표지 제목
        if 'CoverTitle' not in styles:
            cover_style = styles.add_style('CoverTitle', WD_STYLE_TYPE.PARAGRAPH)
            cover_style.font.name = '맑은 고딕'
            cover_style.font.size = Pt(28)
            cover_style.font.bold = True
            cover_style.font.color.rgb = RGBColor(0, 32, 96)  # 진한 파랑
        
        # 섹션 제목
        if 'SectionTitle' not in styles:
            section_style = styles.add_style('SectionTitle', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.name = '맑은 고딕'
            section_style.font.size = Pt(16)
            section_style.font.bold = True
            section_style.font.color.rgb = RGBColor(0, 112, 192)
        
        # 소제목
        if 'SubTitle' not in styles:
            subtitle_style = styles.add_style('SubTitle', WD_STYLE_TYPE.PARAGRAPH)
            subtitle_style.font.name = '맑은 고딕'
            subtitle_style.font.size = Pt(12)
            subtitle_style.font.bold = True
    
    def add_cover_page(
        self,
        title: str = "LH 매입임대주택 사업 제안서",
        address: str = "",
        parcel_id: str = "",
        submitted_by: str = "ZeroSite",
        submission_date: str = ""
    ):
        """표지 페이지 추가"""
        # 제안서 제목
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.name = '맑은 고딕'
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 32, 96)
        
        # 공백
        for _ in range(3):
            self.doc.add_paragraph()
        
        # 대상 부지 정보
        info_para = self.doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_para.add_run(f"대상 부지: {address}")
        info_run.font.name = '맑은 고딕'
        info_run.font.size = Pt(14)
        
        if parcel_id:
            parcel_para = self.doc.add_paragraph()
            parcel_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            parcel_run = parcel_para.add_run(f"필지번호: {parcel_id}")
            parcel_run.font.name = '맑은 고딕'
            parcel_run.font.size = Pt(11)
            parcel_run.font.color.rgb = RGBColor(89, 89, 89)
        
        # 공백
        for _ in range(5):
            self.doc.add_paragraph()
        
        # 제출 정보
        submit_para = self.doc.add_paragraph()
        submit_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        submit_run = submit_para.add_run(f"제출: {submitted_by}")
        submit_run.font.name = '맑은 고딕'
        submit_run.font.size = Pt(12)
        
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(submission_date or datetime.now().strftime("%Y년 %m월 %d일"))
        date_run.font.name = '맑은 고딕'
        date_run.font.size = Pt(12)
        
        # 페이지 구분
        self.doc.add_page_break()
    
    def add_section_title(self, title: str, level: int = 1):
        """섹션 제목 추가"""
        para = self.doc.add_paragraph()
        run = para.add_run(title)
        run.font.name = '맑은 고딕'
        
        if level == 1:
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 112, 192)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 112, 192)
        else:
            run.font.size = Pt(12)
            run.font.bold = True
        
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
    
    def add_paragraph(self, text: str, bold: bool = False, indent: bool = False):
        """일반 단락 추가"""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = '맑은 고딕'
        run.font.size = Pt(11)
        
        if bold:
            run.font.bold = True
        
        if indent:
            para.paragraph_format.left_indent = Inches(0.25)
    
    def add_bullet_list(self, items: List[str]):
        """불릿 리스트 추가"""
        for item in items:
            para = self.doc.add_paragraph(item, style='List Bullet')
            para.runs[0].font.name = '맑은 고딕'
            para.runs[0].font.size = Pt(11)
    
    def add_numbered_list(self, items: List[str]):
        """번호 리스트 추가"""
        for item in items:
            para = self.doc.add_paragraph(item, style='List Number')
            para.runs[0].font.name = '맑은 고딕'
            para.runs[0].font.size = Pt(11)
    
    def add_table(
        self,
        headers: List[str],
        rows: List[List[str]],
        col_widths: Optional[List[float]] = None
    ):
        """표 추가"""
        table = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # 헤더
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # 헤더 스타일
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = '맑은 고딕'
                    run.font.size = Pt(11)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 데이터 행
        for i, row_data in enumerate(rows, start=1):
            row_cells = table.rows[i].cells
            for j, cell_data in enumerate(row_data):
                row_cells[j].text = str(cell_data)
                for paragraph in row_cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '맑은 고딕'
                        run.font.size = Pt(10)
        
        # 열 너비 설정
        if col_widths:
            for i, width in enumerate(col_widths):
                for row in table.rows:
                    row.cells[i].width = Inches(width)
    
    def add_key_value_table(self, data: Dict[str, Any]):
        """키-값 테이블 추가"""
        table = self.doc.add_table(rows=len(data), cols=2)
        table.style = 'Light Grid Accent 1'
        
        for i, (key, value) in enumerate(data.items()):
            # 키 셀
            key_cell = table.rows[i].cells[0]
            key_cell.text = str(key)
            for paragraph in key_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = '맑은 고딕'
                    run.font.size = Pt(10)
            
            # 값 셀
            value_cell = table.rows[i].cells[1]
            value_cell.text = str(value)
            for paragraph in value_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '맑은 고딕'
                    run.font.size = Pt(10)
        
        # 열 너비
        table.rows[0].cells[0].width = Inches(2.0)
        table.rows[0].cells[1].width = Inches(4.0)
    
    def add_financial_summary_table(
        self,
        land_value: int,
        construction_cost: int,
        total_cost: int,
        total_revenue: int,
        npv: int,
        irr: float
    ):
        """재무 요약 표 추가"""
        self.add_section_title("재무 개요", level=2)
        
        data = {
            "토지매입비": f"₩{land_value:,}",
            "건축비": f"₩{construction_cost:,}",
            "총 사업비": f"₩{total_cost:,}",
            "예상 수익": f"₩{total_revenue:,}",
            "순현재가치 (NPV)": f"₩{npv:,}",
            "내부수익률 (IRR)": f"{irr:.2f}%"
        }
        
        self.add_key_value_table(data)
    
    def add_lh_scorecard_table(
        self,
        section_scores: Dict[str, float],
        total_score: float,
        judgement: str,
        grade: str
    ):
        """LH 점수표 추가"""
        self.add_section_title("LH 종합 평가 결과", level=2)
        
        headers = ["평가 항목", "배점", "득점", "득점률"]
        
        section_names = {
            "A": "정책·유형 적합성",
            "B": "입지·환경 평가",
            "C": "건축 가능성",
            "D": "가격·매입 적정성",
            "E": "사업성"
        }
        
        max_scores = {"A": 25, "B": 20, "C": 20, "D": 15, "E": 20}
        
        rows = []
        for section_id, section_name in section_names.items():
            score = section_scores.get(section_id, 0.0)
            max_score = max_scores[section_id]
            score_pct = (score / max_score * 100) if max_score > 0 else 0
            rows.append([
                section_name,
                f"{max_score}점",
                f"{score:.1f}점",
                f"{score_pct:.1f}%"
            ])
        
        # 합계
        rows.append([
            "총점",
            "100점",
            f"{total_score:.1f}점",
            f"{total_score:.1f}%"
        ])
        
        self.add_table(headers, rows, col_widths=[3.0, 1.0, 1.0, 1.0])
        
        # 판정 결과
        self.doc.add_paragraph()
        result_para = self.doc.add_paragraph()
        result_run = result_para.add_run(f"최종 판정: {judgement} (등급: {grade})")
        result_run.font.name = '맑은 고딕'
        result_run.font.size = Pt(12)
        result_run.font.bold = True
        
        if judgement == "GO":
            result_run.font.color.rgb = RGBColor(0, 128, 0)  # 녹색
        elif judgement == "CONDITIONAL":
            result_run.font.color.rgb = RGBColor(255, 165, 0)  # 주황색
        else:
            result_run.font.color.rgb = RGBColor(255, 0, 0)  # 빨강
    
    def add_capacity_table(
        self,
        legal_units: int,
        incentive_units: int,
        legal_gfa: float,
        incentive_gfa: float,
        legal_far: float,
        incentive_far: float
    ):
        """건축 규모 표 추가"""
        self.add_section_title("건축 가능 규모", level=2)
        
        headers = ["구분", "법정 용적률", "인센티브 용적률"]
        rows = [
            ["용적률", f"{legal_far:.1f}%", f"{incentive_far:.1f}%"],
            ["연면적", f"{legal_gfa:,.1f}m²", f"{incentive_gfa:,.1f}m²"],
            ["세대수", f"{legal_units}세대", f"{incentive_units}세대"]
        ]
        
        self.add_table(headers, rows, col_widths=[2.0, 2.0, 2.0])
    
    def save(self, file_path: str):
        """문서 저장"""
        # 디렉토리 생성
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        # 저장
        self.doc.save(file_path)
        print(f"✓ Word document saved: {file_path}")
        
        return file_path


__all__ = ['LHDocumentBuilder']
